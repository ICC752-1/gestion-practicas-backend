"""Servicio de negocio para cartas de presentacion automaticas."""

from dataclasses import dataclass
from datetime import UTC, datetime
import logging
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from uuid import uuid4

from docx import Document as WordDocument
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage, Listing, RichText
from fastapi import HTTPException, status

from app.core.config import config
from app.modules.auth.models.user_model import User
from app.modules.auth.utils.enrollment import build_student_enrollment
from app.modules.notifications.models.notification_model import (
    Notification,
    NotificationEventTypeEnum,
    NotificationStatusEnum,
)
from app.modules.notifications.services.notification_service import (
    NotificationService,
)
from app.modules.notifications.utils.notification_event_helpers import (
    _build_email_body,
)
from app.modules.presentation_letters.models.presentation_letter_model import (
    PresentationLetter,
    PresentationLetterTemplate,
)
from app.modules.presentation_letters.repositories.presentation_letter_repository import (
    PresentationLetterRepository,
)
from app.modules.presentation_letters.schemas.presentation_letter_schema import (
    PRACTICE_TYPE_VALUES,
    PresentationLetterGenerateRequest,
    PresentationLetterTemplateUpdateRequest,
)


logger = logging.getLogger(__name__)

STUDENT_ROLE = "Estudiante"
DIRECTOR_ROLE = "Director de carrera"
ADMIN_READ_ROLES = {
    "Encargado de practica",
    "Director de carrera",
    "Secretaria de Carrera",
}

MONTHS_ES = {
    1: "Enero",
    2: "Febrero",
    3: "Marzo",
    4: "Abril",
    5: "Mayo",
    6: "Junio",
    7: "Julio",
    8: "Agosto",
    9: "Septiembre",
    10: "Octubre",
    11: "Noviembre",
    12: "Diciembre",
}

LIBREOFFICE_TIMEOUT_SECONDS = 60
SIGNATURE_IMAGE_MAX_BYTES = 2 * 1024 * 1024
SIGNATURE_IMAGE_EXTENSIONS = {
    "image/png": "png",
    "image/jpeg": "jpg",
}
RICH_TEXT_CONTEXT_FIELDS = (
    "base_intro",
    "student_presentation",
    "practice_description",
    "minimum_hours_clause",
    "insurance_clause",
    "closing_text",
)


@dataclass(frozen=True)
class PresentationLetterDownload:
    """Datos necesarios para entregar una carta generada."""

    letter: PresentationLetter
    path: Path


@dataclass(frozen=True)
class PresentationLetterSignatureImage:
    """Datos necesarios para entregar la imagen de firma administrada."""

    path: Path
    media_type: str


@dataclass(frozen=True)
class PresentationLetterPreviewTemplate:
    """Plantilla temporal usada para renderizar cambios sin persistirlos."""

    practice_type: str
    title: str
    subtitle: str
    base_intro: str
    student_presentation_template: str
    practice_description: str
    minimum_hours: int
    minimum_hours_clause: str
    learning_outcomes: list[str]
    insurance_clause: str
    closing_text: str
    signature_name: str
    signature_role: str
    signature_institution: str
    signature_image_path: str | None


@dataclass(frozen=True)
class PresentationLetterPreviewStudent:
    """Datos representativos para una previsualizacion administrativa."""

    first_name: str = "Camila"
    last_name: str = "Rojas Soto"
    enrollment: str = "12345678924"
    rut: str = "12345678-9"


def _now() -> datetime:
    return datetime.now(UTC).replace(tzinfo=None)


def _role_names(user: User) -> set[str]:
    return {user_role.role.name for user_role in user.roles}


class PresentationLetterService:
    """Aplica reglas de plantillas, generacion, envio y descarga."""

    def __init__(
        self,
        repository: PresentationLetterRepository,
        app_config: type | object = config,
        notification_service: NotificationService | None = None,
    ) -> None:
        self.repository = repository
        self.notification_service = notification_service
        self.storage_root = Path(
            getattr(
                app_config,
                "PRESENTATION_LETTER_STORAGE_DIR",
                "storage/presentation_letters",
            )
        )
        self.docx_template_path = Path(
            getattr(
                app_config,
                "PRESENTATION_LETTER_DOCX_TEMPLATE_PATH",
                "app/core/assets/presentation_letter_template.docx",
            )
        )
        self.libreoffice_binary = getattr(
            app_config,
            "LIBREOFFICE_BINARY",
            "libreoffice",
        )
        self.libreoffice_timeout_seconds = getattr(
            app_config,
            "LIBREOFFICE_TIMEOUT_SECONDS",
            LIBREOFFICE_TIMEOUT_SECONDS,
        )

    async def list_templates(
        self,
        *,
        actor: User,
    ) -> list[PresentationLetterTemplate]:
        """Lista plantillas activas para roles administrativos."""

        self._require_template_reader(actor)
        return await self.repository.list_templates()

    async def get_template(
        self,
        *,
        practice_type: str,
        actor: User,
    ) -> PresentationLetterTemplate:
        """Obtiene la plantilla activa para un tipo de practica."""

        self._require_template_reader(actor)
        return await self._get_template_or_404(practice_type)

    async def update_template(
        self,
        *,
        practice_type: str,
        payload: PresentationLetterTemplateUpdateRequest,
        actor: User,
    ) -> PresentationLetterTemplate:
        """Edita una plantilla. Solo Director de carrera."""

        self._require_director(actor)
        self._validate_practice_type(practice_type)
        template = await self.repository.get_active_template(practice_type)
        timestamp = _now()

        if template is None:
            template = PresentationLetterTemplate(
                practice_type=practice_type,
                created_by=actor.id,
                created_at=timestamp,
            )

        template.title = payload.title
        template.subtitle = payload.subtitle
        template.base_intro = payload.base_intro
        template.student_presentation_template = payload.student_presentation_template
        template.practice_description = payload.practice_description
        template.minimum_hours = payload.minimum_hours
        template.minimum_hours_clause = payload.minimum_hours_clause
        template.learning_outcomes = payload.learning_outcomes
        template.insurance_clause = payload.insurance_clause
        template.closing_text = payload.closing_text
        template.signature_name = payload.signature_name
        template.signature_role = payload.signature_role
        template.signature_institution = payload.signature_institution
        template.is_active = payload.is_active
        template.updated_by = actor.id
        template.updated_at = timestamp

        return await self.repository.save_template(template)

    async def preview_template(
        self,
        *,
        practice_type: str,
        payload: PresentationLetterTemplateUpdateRequest,
        actor: User,
    ) -> bytes:
        """Renderiza la edicion actual como PDF sin guardar la plantilla."""

        self._require_template_reader(actor)
        stored_template = await self._get_template_or_404(practice_type)
        preview_template = PresentationLetterPreviewTemplate(
            practice_type=practice_type,
            title=payload.title,
            subtitle=payload.subtitle,
            base_intro=payload.base_intro,
            student_presentation_template=payload.student_presentation_template,
            practice_description=payload.practice_description,
            minimum_hours=payload.minimum_hours,
            minimum_hours_clause=payload.minimum_hours_clause,
            learning_outcomes=payload.learning_outcomes,
            insurance_clause=payload.insurance_clause,
            closing_text=payload.closing_text,
            signature_name=payload.signature_name,
            signature_role=payload.signature_role,
            signature_institution=payload.signature_institution,
            signature_image_path=getattr(
                stored_template,
                "signature_image_path",
                None,
            ),
        )

        return self._render_pdf(
            template=preview_template,
            student=PresentationLetterPreviewStudent(),
            generated_at=_now(),
        )

    async def update_template_signature_image(
        self,
        *,
        practice_type: str,
        file_name: str,
        content_type: str | None,
        content: bytes,
        actor: User,
    ) -> PresentationLetterTemplate:
        """Reemplaza la imagen de firma de una plantilla. Solo Director."""

        self._require_director(actor)
        template = await self._get_template_or_404(practice_type)
        media_type, extension = self._validate_signature_image(
            file_name=file_name,
            content_type=content_type,
            content=content,
        )
        old_storage_key = template.signature_image_path
        storage_key = (
            f"signatures/{_slugify(practice_type)}-{uuid4().hex}.{extension}"
        )
        self._write_file(storage_key, content)

        template.signature_image_path = storage_key
        template.updated_by = actor.id
        template.updated_at = _now()

        try:
            saved = await self.repository.save_template(template)
        except Exception:
            self._resolve_storage_key(storage_key).unlink(missing_ok=True)
            raise

        if old_storage_key and old_storage_key != storage_key:
            self._resolve_storage_key(old_storage_key).unlink(missing_ok=True)

        logger.info(
            "Presentation letter signature image updated",
            extra={
                "actor_id": actor.id,
                "practice_type": practice_type,
                "media_type": media_type,
            },
        )

        return saved

    async def remove_template_signature_image(
        self,
        *,
        practice_type: str,
        actor: User,
    ) -> PresentationLetterTemplate:
        """Elimina la imagen de firma administrada de una plantilla."""

        self._require_director(actor)
        template = await self._get_template_or_404(practice_type)
        old_storage_key = template.signature_image_path
        template.signature_image_path = None
        template.updated_by = actor.id
        template.updated_at = _now()
        saved = await self.repository.save_template(template)

        if old_storage_key:
            self._resolve_storage_key(old_storage_key).unlink(missing_ok=True)

        return saved

    async def prepare_signature_image(
        self,
        *,
        practice_type: str,
        actor: User,
    ) -> PresentationLetterSignatureImage:
        """Prepara la imagen de firma para vista administrativa."""

        self._require_template_reader(actor)
        template = await self._get_template_or_404(practice_type)
        if not template.signature_image_path:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="La plantilla no tiene una imagen de firma configurada.",
            )

        path = self._resolve_storage_key(template.signature_image_path)
        if not path.is_file():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No se encontró la imagen de firma configurada.",
            )

        return PresentationLetterSignatureImage(
            path=path,
            media_type=self._media_type_for_signature_path(path),
        )

    async def generate_letter(
        self,
        *,
        actor: User,
        payload: PresentationLetterGenerateRequest,
    ) -> PresentationLetter:
        """Genera automaticamente una carta PDF y registra su envio."""

        self._require_student(actor)
        template = await self._get_template_or_404(payload.practice_type)
        generated_at = _now()
        storage_key = self._build_storage_key(
            student_id=actor.id,
            practice_type=payload.practice_type,
            generated_at=generated_at,
        )
        file_name = self._build_file_name(actor, payload.practice_type)
        pdf_bytes = self._render_pdf(
            template=template,
            student=actor,
            generated_at=generated_at,
        )
        stored_path = self._write_file(storage_key, pdf_bytes)

        letter = PresentationLetter(
            student_id=actor.id,
            practice_type=payload.practice_type,
            template_id=template.id,
            generated_file_name=file_name,
            generated_file_path=storage_key,
            recipient_email=actor.email,
            created_at=generated_at,
            updated_at=generated_at,
        )

        try:
            created = await self.repository.create_letter(letter)
        except Exception:
            stored_path.unlink(missing_ok=True)
            raise

        dispatch_ok = await self._notify_student_letter_generated(created)
        if dispatch_ok:
            created.sent_at = _now()
            created.updated_at = created.sent_at
            created = await self.repository.save_letter(created)

        return created

    async def list_my_letters(
        self,
        *,
        actor: User,
    ) -> list[PresentationLetter]:
        """Lista cartas generadas del estudiante autenticado."""

        self._require_student(actor)
        return await self.repository.list_letters_for_student(actor.id)

    async def prepare_download(
        self,
        *,
        letter_id: int,
        actor: User,
    ) -> PresentationLetterDownload:
        """Prepara descarga autenticada de una carta generada."""

        letter = await self.repository.get_letter_by_id(letter_id)
        if letter is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Carta de presentacion no encontrada",
            )

        if not (self._is_admin(actor) or letter.student_id == actor.id):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="No tienes permisos para descargar esta carta",
            )

        file_path = self._resolve_storage_key(letter.generated_file_path)
        if not file_path.is_file():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Archivo de carta no encontrado",
            )

        letter.downloaded_at = _now()
        letter.updated_at = letter.downloaded_at
        updated = await self.repository.save_letter(letter)

        return PresentationLetterDownload(letter=updated, path=file_path)

    def _require_student(self, user: User) -> None:
        if STUDENT_ROLE not in _role_names(user):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Solo estudiantes pueden generar cartas de presentacion",
            )

    def _require_director(self, user: User) -> None:
        if DIRECTOR_ROLE not in _role_names(user):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Solo el Director de carrera puede editar plantillas",
            )

    def _require_template_reader(self, user: User) -> None:
        if not (_role_names(user) & ADMIN_READ_ROLES):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="No tienes permisos para consultar plantillas",
            )

    def _is_admin(self, user: User) -> bool:
        return bool(_role_names(user) & ADMIN_READ_ROLES)

    @staticmethod
    def _validate_practice_type(practice_type: str) -> None:
        if practice_type not in PRACTICE_TYPE_VALUES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "code": "invalid_practice_type",
                    "message": "Tipo de practica no soportado para carta de presentacion",
                    "allowed_values": list(PRACTICE_TYPE_VALUES),
                },
            )

    async def _get_template_or_404(
        self,
        practice_type: str,
    ) -> PresentationLetterTemplate:
        self._validate_practice_type(practice_type)
        template = await self.repository.get_active_template(practice_type)
        if template is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "code": "presentation_letter_template_not_found",
                    "message": (
                        "No existe una plantilla activa para el tipo de practica "
                        "seleccionado."
                    ),
                    "practice_type": practice_type,
                },
            )

        return template

    def _render_pdf(
        self,
        *,
        template: PresentationLetterTemplate | PresentationLetterPreviewTemplate,
        student: User | PresentationLetterPreviewStudent,
        generated_at: datetime,
    ) -> bytes:
        return self._render_docx_template_to_pdf(
            self._build_letter_document(
                template=template,
                student=student,
                generated_at=generated_at,
            )
        )

    def _build_letter_document(
        self,
        *,
        template: PresentationLetterTemplate | PresentationLetterPreviewTemplate,
        student: User | PresentationLetterPreviewStudent,
        generated_at: datetime,
    ) -> dict[str, object]:
        variables = self._build_variables(
            template=template,
            student=student,
            generated_at=generated_at,
        )
        rendered_learning_outcomes = [
            self._render_template_text(outcome, variables)
            for outcome in template.learning_outcomes
        ]
        learning_text = "\n".join(f"• {outcome}" for outcome in rendered_learning_outcomes)
        variables["learning_outcomes"] = learning_text
        paragraph_templates = [
            template.base_intro,
            template.student_presentation_template,
            template.practice_description,
            template.minimum_hours_clause,
        ]

        return {
            "date": f"Temuco, {generated_at.day} de {MONTHS_ES[generated_at.month]} del {generated_at.year}",
            "title": self._render_template_text(template.title, variables).upper(),
            "subtitle": self._render_template_text(template.subtitle, variables),
            "greeting": "A quien corresponda:",
            "paragraph_templates": paragraph_templates,
            "paragraphs": [
                self._render_template_text(text, variables)
                for text in paragraph_templates
            ],
            "learning_outcomes": rendered_learning_outcomes,
            "variables": variables,
            "insurance_clause_template": template.insurance_clause,
            "insurance_clause": self._render_template_text(
                template.insurance_clause,
                variables,
            ),
            "closing_text_template": template.closing_text,
            "closing_text": self._render_template_text(template.closing_text, variables),
            "signature_image_path": (
                self._resolve_storage_key(signature_image_path)
                if (signature_image_path := getattr(template, "signature_image_path", None))
                else None
            ),
            "signature_name": template.signature_name,
            "signature_role": template.signature_role,
            "signature_institution": template.signature_institution,
        }

    @staticmethod
    def _build_variables(
        *,
        template: PresentationLetterTemplate | PresentationLetterPreviewTemplate,
        student: User | PresentationLetterPreviewStudent,
        generated_at: datetime,
    ) -> dict[str, str]:
        student_name = f"{student.first_name} {student.last_name}".strip()
        identifier = build_student_enrollment(student) or student.rut
        return {
            "student_name": student_name,
            "student_identifier": identifier,
            "practice_type": template.practice_type,
            "current_date": (
                f"{generated_at.day} de {MONTHS_ES[generated_at.month]} "
                f"del {generated_at.year}"
            ),
            "minimum_hours": str(template.minimum_hours),
        }

    @staticmethod
    def _render_template_text(text: str, variables: dict[str, str]) -> str:
        rendered = text
        for key, value in variables.items():
            rendered = rendered.replace(f"{{{{{key}}}}}", value)
        return rendered

    def _render_docx_template_to_pdf(self, document: dict[str, object]) -> bytes:
        template_path = self._resolve_docx_template_path()
        with tempfile.TemporaryDirectory(prefix="presentation-letter-") as temp_dir:
            output_dir = Path(temp_dir)
            rich_text_template = output_dir / "plantilla-carta.docx"
            rendered_docx = output_dir / "carta-presentacion.docx"

            self._prepare_rich_text_template(
                source=template_path,
                destination=rich_text_template,
            )
            docx_template = DocxTemplate(str(rich_text_template))
            docx_template.render(self._build_docx_context(document, docx_template))
            docx_template.save(str(rendered_docx))

            return self._convert_docx_to_pdf(rendered_docx, output_dir)

    @staticmethod
    def _prepare_rich_text_template(*, source: Path, destination: Path) -> None:
        """Convierte marcadores de texto editable en marcadores RichText."""

        document = WordDocument(str(source))
        placeholders = {
            f"{{{{ {field} }}}}": f"{{{{r {field} }}}}"
            for field in RICH_TEXT_CONTEXT_FIELDS
        }

        for paragraph in document.paragraphs:
            replacement = placeholders.get(paragraph.text.strip())
            if replacement is None:
                continue
            paragraph.clear()
            paragraph.add_run(replacement)

        document.save(str(destination))

    def _resolve_docx_template_path(self) -> Path:
        template_path = self.docx_template_path.resolve()
        if not template_path.is_file():
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=(
                    "No se encontro la plantilla DOCX institucional para cartas "
                    "de presentacion."
                ),
            )

        return template_path

    @staticmethod
    def _build_docx_context(
        document: dict[str, object],
        docx_template: DocxTemplate | None = None,
    ) -> dict[str, object]:
        paragraph_templates = list(document["paragraph_templates"])
        variables = dict(document["variables"])
        learning_outcomes_text = "\n".join(
            f"• {outcome}" for outcome in document["learning_outcomes"]
        )
        signature_image_path = document.get("signature_image_path")
        signature_image = (
            InlineImage(docx_template, str(signature_image_path), width=Mm(42))
            if signature_image_path and docx_template is not None
            else ""
        )

        return {
            "date": document["date"],
            "title": document["title"],
            "subtitle": document["subtitle"],
            "greeting": document["greeting"],
            "base_intro": PresentationLetterService._render_template_rich_text(
                paragraph_templates[0],
                variables,
            ),
            "student_presentation": PresentationLetterService._render_template_rich_text(
                paragraph_templates[1],
                variables,
            ),
            "practice_description": PresentationLetterService._render_template_rich_text(
                paragraph_templates[2],
                variables,
            ),
            "minimum_hours_clause": PresentationLetterService._render_template_rich_text(
                paragraph_templates[3],
                variables,
            ),
            "learning_outcomes_text": Listing(learning_outcomes_text),
            "insurance_clause": PresentationLetterService._render_template_rich_text(
                str(document["insurance_clause_template"]),
                variables,
            ),
            "closing_text": PresentationLetterService._render_template_rich_text(
                str(document["closing_text_template"]),
                variables,
            ),
            "signature_image": signature_image,
            "signature_name": document["signature_name"],
            "signature_role": document["signature_role"],
            "signature_institution": document["signature_institution"],
        }

    @staticmethod
    def _render_template_rich_text(
        text: str,
        variables: dict[str, str],
    ) -> RichText:
        """Renderiza variables como segmentos en negrita dentro de un párrafo."""

        rich_text = RichText()
        for part in re.split(r"({{[a-z_]+}})", text):
            match = re.fullmatch(r"{{([a-z_]+)}}", part)
            if match and match.group(1) in variables:
                rich_text.add(variables[match.group(1)], bold=True)
            elif part:
                rich_text.add(part)
        return rich_text

    @staticmethod
    def _validate_signature_image(
        *,
        file_name: str,
        content_type: str | None,
        content: bytes,
    ) -> tuple[str, str]:
        if not content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="La imagen de firma está vacía.",
            )

        if len(content) > SIGNATURE_IMAGE_MAX_BYTES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="La imagen de firma no puede superar 2 MB.",
            )

        detected_media_type: str | None = None
        if content.startswith(b"\x89PNG\r\n\x1a\n"):
            detected_media_type = "image/png"
        elif content.startswith(b"\xff\xd8\xff"):
            detected_media_type = "image/jpeg"

        if detected_media_type is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Sube una imagen de firma válida en formato PNG o JPG.",
            )

        normalized_content_type = (content_type or "").split(";")[0].lower()
        if normalized_content_type and normalized_content_type != detected_media_type:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="El tipo de archivo no coincide con la imagen enviada.",
            )

        extension = SIGNATURE_IMAGE_EXTENSIONS[detected_media_type]
        suffix = Path(file_name or "").suffix.lower().lstrip(".")
        allowed_suffixes = {"png"} if extension == "png" else {"jpg", "jpeg"}
        if suffix and suffix not in allowed_suffixes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="La extensión de la firma debe ser PNG, JPG o JPEG.",
            )

        return detected_media_type, extension

    @staticmethod
    def _media_type_for_signature_path(path: Path) -> str:
        if path.suffix.lower() == ".png":
            return "image/png"
        return "image/jpeg"

    def _convert_docx_to_pdf(self, docx_path: Path, output_dir: Path) -> bytes:
        libreoffice = self._resolve_libreoffice_binary()
        profile_dir = output_dir / "libreoffice-profile"
        profile_dir.mkdir(parents=True, exist_ok=True)
        home_dir = output_dir / "home"
        runtime_dir = output_dir / "xdg-runtime"
        home_dir.mkdir(parents=True, exist_ok=True)
        runtime_dir.mkdir(parents=True, exist_ok=True)
        runtime_dir.chmod(0o700)
        env = os.environ.copy()
        env.update(
            {
                "HOME": str(home_dir),
                "XDG_RUNTIME_DIR": str(runtime_dir),
                "SAL_USE_VCLPLUGIN": "svp",
            }
        )
        command = [
            libreoffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]

        try:
            result = subprocess.run(
                command,
                capture_output=True,
                check=False,
                env=env,
                text=True,
                timeout=self.libreoffice_timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="La conversion de la carta a PDF excedio el tiempo maximo.",
            ) from exc

        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        if result.returncode != 0 or not pdf_path.is_file():
            logger.error(
                "LibreOffice no pudo convertir carta DOCX a PDF. stdout=%s stderr=%s",
                result.stdout,
                result.stderr,
            )
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=_libreoffice_error_detail(result.stderr),
            )

        return pdf_path.read_bytes()

    def _resolve_libreoffice_binary(self) -> str:
        resolved = shutil.which(self.libreoffice_binary)
        if resolved is not None:
            return resolved

        binary_path = Path(self.libreoffice_binary)
        if binary_path.is_file():
            return str(binary_path)

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="LibreOffice no esta disponible para generar la carta PDF.",
        )

    @staticmethod
    def _build_file_name(student: User, practice_type: str) -> str:
        student_name = f"{student.first_name}-{student.last_name}".strip("-")
        slug = _slugify(f"carta-presentacion-{student_name}-{practice_type}")
        return f"{slug}.pdf"

    @staticmethod
    def _build_storage_key(
        *,
        student_id: int,
        practice_type: str,
        generated_at: datetime,
    ) -> str:
        practice_slug = _slugify(practice_type)
        date_slug = generated_at.strftime("%Y%m%d-%H%M%S")
        return f"{student_id}/{practice_slug}/{date_slug}-{uuid4().hex}.pdf"

    def _write_file(self, storage_key: str, content: bytes) -> Path:
        file_path = self._resolve_storage_key(storage_key)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(content)
        return file_path

    def _resolve_storage_key(self, storage_key: str) -> Path:
        root = self.storage_root.resolve()
        file_path = (root / storage_key).resolve()
        if root != file_path and root not in file_path.parents:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Ruta de archivo invalida",
            )

        return file_path

    async def _notify_student_letter_generated(
        self,
        letter: PresentationLetter,
    ) -> bool:
        if self.notification_service is None:
            return False

        notification = Notification(
            recipient_user_id=letter.student_id,
            recipient_email=letter.recipient_email,
            event_type=NotificationEventTypeEnum.custom,
            subject="Carta de presentación generada",
            content=_build_email_body(
                title="Carta de presentación generada",
                intro=(
                    "Su carta de presentación fue generada automáticamente. "
                    "Puede descargarla desde la plataforma con su sesión activa."
                ),
                details=[
                    ("Tipo de práctica", letter.practice_type),
                    ("Archivo", letter.generated_file_name),
                    ("ID carta", letter.id),
                ],
                action_label="Descargar carta",
            ),
            status=NotificationStatusEnum.simulated,
            payload={
                "event": "presentation_letter_generated",
                "presentation_letter_id": letter.id,
                "practice_type": letter.practice_type,
                "generated_file_name": letter.generated_file_name,
            },
        )

        try:
            await self.notification_service.create_and_dispatch(notification)
            return True
        except Exception as exc:
            logger.warning(
                "No se pudo registrar envio de carta id=%s: %s",
                letter.id,
                exc,
                exc_info=True,
            )
            return False


def _slugify(value: str) -> str:
    normalized = value.lower()
    replacements = {
        "á": "a",
        "é": "e",
        "í": "i",
        "ó": "o",
        "ú": "u",
        "ñ": "n",
    }
    for source, target in replacements.items():
        normalized = normalized.replace(source, target)
    normalized = re.sub(r"[^a-z0-9]+", "-", normalized).strip("-")
    return normalized or "carta-presentacion"


def _libreoffice_error_detail(stderr: str) -> str:
    if "source file could not be loaded" in stderr:
        return (
            "LibreOffice no pudo abrir el DOCX. Verifica que este instalado "
            "libreoffice-writer en el entorno donde corre el backend."
        )

    return "No se pudo convertir la carta de presentacion a PDF."
